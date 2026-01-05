'''
This code file is used to save the model outputs to a specified location within the Answers folder.
'''

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from typing import List, Union
import pandas as pd



def set_table_borders(table):
    """
    Set all borders of the table to be visible.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), 'auto')
        tblBorders.append(edge_el)

    tblPr.append(tblBorders)
    


def standardize_speaker_format(data: Union[List[str], List[List[str]]]) -> Union[List[str], List[List[str]]]:
    speaker_map = {
        "Co-Ordinator: ": "[Co-Ordinator]：",
        "Plant: ": "[Plant]：",
        "Monitor Evaluator: ": "[Monitor Evaluator]：",
        "Implementer: ": "[Implementer]：",
    }

    def replace_speaker_prefix(text: str) -> str:
        for old, new in speaker_map.items():
            if text.startswith(old):
                return text.replace(old, new, 1)
        return text

    if isinstance(data[0], list):
        return [[replace_speaker_prefix(text) for text in sublist] for sublist in data]
    else:
        return [replace_speaker_prefix(text) for text in data]



def pre_text(history):
    h0 = standardize_speaker_format(history[0])
    h1 = standardize_speaker_format(history[1])
    return h0, h1, history[2]



def complete_answer(history, name, dir):
    history = pre_text(history)

    doc0 = Document()
    doc1 = Document()
    doc2 = Document()
    table1 = doc1.add_table(rows = len(history[1]) + 1, cols = 1)
    table2 = doc2.add_table(rows = len(history[2]) + 1, cols = 1)

    table1.cell(0, 0).text = name
    table2.cell(0, 0).text = name

    for j in [1, 2]:
        for idx, item in enumerate(history[j], start = 1):
            cell = table1.cell(idx, 0) if j == 1 else table2.cell(idx, 0)
            cell.text = '\n\n=====================================================================\n\n'.join(item) \
                if j == 1 else item

    set_table_borders(table1)
    set_table_borders(table2)
    os.makedirs(os.path.join("Answers", "Full_Step_History_Record", dir), exist_ok = True)
    doc1.save(os.path.join("Answers", "Full_Step_History_Record", dir, f"{name}.docx"))
    print("Full_Step_History_Record saved!")
    os.makedirs(os.path.join("Answers", "Answer", dir), exist_ok = True)
    doc2.save(os.path.join("Answers", "Answer", dir, f"{name}.docx"))
    print("Answer saved!")

    his = name + '\n\n' + "\n\n".join(history[0])   
    doc0.add_paragraph(his)
    os.makedirs(os.path.join("Answers", "Full_History_Record", dir), exist_ok = True)
    doc0.save(os.path.join("Answers", "Full_History_Record", dir, f"{name}.docx"))
    print("Full_History_Record saved!")



def write_time(Time, name, FS_lst):
    df = pd.DataFrame([FS_lst, Time])
    os.makedirs("Answers\Time", exist_ok = True)
    filename = os.path.join("Answers\Time", name + '.xlsx')
    df.to_excel(filename, index = False, header = False)
    print(name, "- Time saved!")



def write_single_answer(answer, name, dir):
    doc = Document()
    table = doc.add_table(rows = len(answer) + 1, cols = 1)

    table.cell(0, 0).text = name

    for idx, item in enumerate(answer, start = 1):
        cell = table.cell(idx, 0)
        cell.text = item

    set_table_borders(table)
    os.makedirs(os.path.join("Answers", "Answer", dir), exist_ok = True)
    doc.save(os.path.join("Answers", "Answer", dir, f"{name}.docx"))
    print("Answer saved!")