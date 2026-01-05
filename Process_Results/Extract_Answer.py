import os
import pandas as pd
import numpy as np
from docx import Document
from bert_score import BERTScorer
from scipy.stats import pearsonr
import math
from Extract_Scores import get_original_scores, sort_scores, check, pairwise_average, steps_dimensions_max_wo_integrity,\
        steps_dimensions_wo_integrity, steps_dimensions_max_with_integrity, steps_dimensions_with_integrity, \
        human_allocation_A, human_allocation_B, ScoreName_lst, model_name
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Border, Side, Font
from scipy.stats import wilcoxon
from draw import write_integrity, ini, average_FS
import jieba
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
import matplotlib.pyplot as plt


def sort_scores(scores_wo_sort: list) -> list:
    '''
    Sort a list in the PScore format by response ID.
    '''
    def sort_key(item):
        part1, part2 = item[0].split('_') if isinstance(item, list) else item.split('_')
        condition = part1[0]
        num1 = int(part1[1:])
        fs_num = int(part2[2:])
        # human_id = int(item[1][1:])
        return (condition, fs_num, num1) # (fs_num, human_id, num1)
    return sorted(scores_wo_sort, key = sort_key)



def list_folders(path):
    return [f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))]



def list_word_files(folder_path: str) -> list: 
    '''
    Return a list of all Word file names (without paths) in the specified folder
    '''
    word_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    return word_files



def read_table_from_docx(filepath):
    """
    Read the content of the table at a specified index in a Word file
    and return it in list form: [ans_id, {'Step-1': ***, 'Step-2': ***, .....}]
    """
    doc = Document(filepath)
    table = doc.tables[0]
    data = []
    ans_dic = {}
    for i, row in enumerate(table.rows):
        row_text = [cell.text.strip() for cell in row.cells]
        if i == 0: data.append(row_text[0])
        else: ans_dic['Step-' + str(i)] = row_text[0]
    data.append(ans_dic)
    return data



def get_original_answers(file_path: str) -> list:
    '''
    This function traverses all Word rating sheet files and all worksheets within the Human Evaluation directory, and extracts the scores for each evaluation dimension at every step.

    Inputs: Relative path to the folder, e.g., 'Human Evaluation'

    Outputs: A list in the format of PAnswer 
    '''
    PAnswer = [] # len == 200
    ans_ID_lst = list_folders(file_path)
    
    for ans_id in ans_ID_lst:
        all_ans_files = list_word_files(os.path.join(file_path, ans_id))
        assert len(all_ans_files) == 10, ans_id + 'There are fewer than ten responses!'
        for ans_file in all_ans_files:
            PAnswer.append(read_table_from_docx(os.path.join(file_path, ans_id, ans_file)))

    return PAnswer



def extract_step_answers(PAnswer, ans_id, step):
    '''
    Extract the answers of a specified step from the 20 responses corresponding to a given model ID.
    '''
    answers = []
    for condition in ['A', 'B']:
        answers.extend([[ans[0], ans[1][step]] for ans in PAnswer if ans[0].startswith(condition + ans_id)])

    return answers



def write_answers_to_docx(answers, save_path):
    """
    Write the answers into a Word file, with each answer placed in a separate 2×1 table:

    Row 1: the ID

    Row 2: the response content
    """
    doc = Document()
    
    for ans_id, content in answers:
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid" # Add borders to a table
        table.cell(0, 0).text = ans_id # First row: write the ID
        table.cell(1, 0).text = str(content) # Second row: write the answer content
        doc.add_paragraph("")

    doc.save(save_path)
    print(f"✅ Saved to: {save_path}")



def list_challenge_and_solution(text):
    '''
    List the text content of Step-1 or Step-3, organized by scenario/challenge ID.
    '''
    lines, content = text.split('\n'), []
    start, p = 1, 1
    if p:
        for line in lines:
            if line.startswith(str(start)):
                content.append(line.lstrip()[2:])
                start += 1
            else: p = 0

    return content



def compute_step1_or_3_selfbleu(sentences):
    """
    step3_answers: a list of answers, containing 8 items corresponding to 8 scenarios/challenges.
    Returns: The average similarity between the answers
    """
    scores, smooth = [], SmoothingFunction().method1
    grams = [list(jieba.cut(s)) for s in sentences]
    for i in range(len(sentences)):
   
        hypothesis = grams[i]
        references = [grams[j] for j in range(len(grams)) if j != i]

        score = sentence_bleu(references, hypothesis, smoothing_function=smooth, weights=(0.8, 0.2))
        # weights=(0.8, 0.2), (0.7, 0.3), (0.6, 0.25, 0.1, 0.05), (0.25, 0.25, 0.25, 0.25)
        scores.append(score)
    return sum(scores) / len(scores)



def cal_step_similarity(PAnswer, AScore_T, AScore_S, step):
    '''
    step: 'Step-1' or 'Step-3'
    '''
    step_answer = []
    plot1 = [ans[1][step]['Flexibility'] for ans in AScore_T]
    plot2 = [ans[1][step]['Flexibility'] for ans in AScore_S]
    for ans in ans_lst:
        step_answer.extend([[an[0], list_challenge_and_solution(an[1][step])] for an in PAnswer if an[0] == ans])
    
    simi = []
    model_ID_lst = ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10']
    for model_id in model_ID_lst:
        for FS in fs_order:
            model_A = [ans for ans in step_answer if ans[0] == 'A' + model_id + '_' + FS][0]
            model_B = [ans for ans in step_answer if ans[0] == 'B' + model_id + '_' + FS][0]
            # Diversity = 1 - Average Self-BLEU
            simi.append([model_A[0], round(1 - compute_step1_or_3_selfbleu(model_A[1]), 4)])
            # print(simi[-1][0], simi[-1][1])
            simi.append([model_B[0], round(1 - compute_step1_or_3_selfbleu(model_B[1]), 4)])
            # print(simi[-1][0], simi[-1][1])

    simi = np.array(simi)
    # simi_T = [ans for ans in simi if ans[0].startswith('A')]
    # simi_S = [ans for ans in simi if ans[0].startswith('B')]
    # draw_table_simi_comparison(simi_T, simi_S, model_ID_lst, plot1, plot2, step, 'Results', 'Similarity_Comparison_' + step + '_selfbleu.xlsx', 0.01)
    # r, p = pearsonr(np.array(simi[:, 1].astype(float)), np.array(simi[:, 2].astype(float)))
    

if __name__ == '__main__':
    PAnswer = get_original_answers('Answers\Answer')
    model_lst = sorted(set([ans[0][:3] for ans in PAnswer]))
    ans_lst = sorted(set([ans[0] for ans in PAnswer]))
    fs_order = [f"FS{i}" for i in range(1, 11)] # Future scenario numbers
    '''
    PAnswer = [
            [ans_id, {'Step-1': ***, 'Step-2': ***, ...}],
            ...
        ]
    '''
    ''' 1. Process the original scores in a double-blind manner, without merging the Step-2 dimensions. '''
    file_path_Team = os.path.join('Human Evaluation', 'DATA_ACC')
    file_path_Single = os.path.join('Human Evaluation', 'DATA_BCC')

    PScore_T = get_original_scores(file_path_Team)
    check(PScore_T, human_allocation_A)
    print()
    PScore_with_average_T = pairwise_average(PScore_T) # Average the scores from two evaluators
    PScore_with_merge_T = write_integrity(PScore_with_average_T) # Merge dimensions
    EScore_T = sort_scores(PScore_with_merge_T) # Team Work Formal Scores
    model_ID_lst_T, FS_lst = ini(EScore_T)
    AScore_T = average_FS(EScore_T, model_ID_lst_T, FS_lst) # Final Average Scores for Team Work

    PScore_S = get_original_scores(file_path_Single)
    check(PScore_S, human_allocation_B)
    PScore_with_average_S = pairwise_average(PScore_S)
    PScore_with_merge_S = write_integrity(PScore_with_average_S)
    EScore_S = sort_scores(PScore_with_merge_S) # Single Agent Formal Scores
    model_ID_lst_S, FS_lst = ini(EScore_S)
    AScore_S = average_FS(EScore_S, model_ID_lst_S, FS_lst)

    ''' Observe llama model performance in Step-5 '''
    # llama_in_step_5 = sort_scores(extract_step_answers(PAnswer, '06', 'Step-5'))
    # write_answers_to_docx(llama_in_step_5,  "D:\Desktop\llama_in_step_5.docx")


    ''' Calculate similarity of Step-3 solutions '''
    cal_step_similarity(PAnswer, AScore_T, AScore_S, 'Step-1')
    cal_step_similarity(PAnswer, AScore_T, AScore_S, 'Step-3')