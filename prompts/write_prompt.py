'''
The code file is used to write FPSP.json, prompts.json, and config_role.json files,
so that relevant fields can be modified in the json files in a timely manner.
'''

from docx import Document
import json
docx_path = 'Prompt_English.docx'
team_role_path = r'prompts\team_role.json'
FPSP_path = r'prompts\FPSP.json'
prompts_path = r'prompts\prompts.json'

def team_role():
    doc = Document(docx_path)
    table = doc.tables[0]
    data = []

    for _, row in enumerate(table.rows[1:], start = 1):
        role = {
            "Agent_name": row.cells[0].text,
            "Role_Speciality": row.cells[1].text,
            "Role_Prompt": row.cells[2].text
        }
        data.append(role)
    
    with open(team_role_path, 'w', encoding = "utf-8") as file:
        json.dump(data, file, indent = 4, ensure_ascii=False)
    
    print(team_role_path, ' file has been successfully written!')
    

def FPSP():
    doc = Document(docx_path)
    table = doc.tables[1]
    data = []
    headers = [cell.text for cell in table.rows[0].cells]
    
    for i, row in enumerate(table.rows[1:], start = 1):
        row_data = {
            'Step_Number': str(i),
            headers[0]: row.cells[0].text,
            headers[1]: row.cells[1].text,
            headers[2]: row.cells[2].text
        }
        data.append(row_data)
    
    with open(FPSP_path, 'w', encoding = "utf-8") as file:
        json.dump(data, file, indent = 4, ensure_ascii=False)
    
    print(FPSP_path, ' file has been successfully written!')


def prompts():

    doc = Document(docx_path)
    table = doc.tables[2]
    data = {}
    headers = [cell.text for cell in table.columns[0].cells]
    
    for i, row in enumerate(table.rows[1:], start = 1):
        data[headers[i]] = row.cells[1].text
    
    with open(prompts_path, 'w', encoding = "utf-8") as file:
        json.dump(data, file, indent = 4, ensure_ascii=False)
    
    print(prompts_path, ' file has been successfully written!')


team_role()
FPSP()
prompts()